#!/usr/bin/env python3
"""
Advanced Markdown to Word conversion with better equation handling
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_advanced_word_document(markdown_file, output_file):
    """Convert Markdown file to Word document with advanced formatting"""
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.space_after = Pt(18)
    
    # Heading 1 style
    heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.space_after = Pt(12)
    heading1_style.space_before = Pt(12)
    
    # Heading 2 style
    heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.space_after = Pt(10)
    heading2_style.space_before = Pt(10)
    
    # Heading 3 style
    heading3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True
    heading3_style.space_after = Pt(8)
    heading3_style.space_before = Pt(8)
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)
    normal_style.space_after = Pt(6)
    
    # Equation style
    equation_style = styles.add_style('CustomEquation', WD_STYLE_TYPE.PARAGRAPH)
    equation_style.font.size = Pt(11)
    equation_style.font.name = 'Times New Roman'
    equation_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_style.space_after = Pt(8)
    equation_style.space_before = Pt(8)
    
    # Table caption style
    caption_style = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.space_after = Pt(12)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle title (first line)
        if i == 0 and line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='CustomTitle')
            i += 1
            continue
        
        # Handle headings
        if line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_paragraph(heading, style='CustomHeading1')
            i += 1
            continue
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_paragraph(heading, style='CustomHeading2')
            i += 1
            continue
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_paragraph(heading, style='CustomHeading3')
            i += 1
            continue
        
        # Handle table captions
        if line.startswith('**Table') and '**' in line:
            caption = line.replace('**', '').strip()
            p = doc.add_paragraph(caption, style='CustomCaption')
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                # Parse table
                table_data = []
                for table_line in table_lines:
                    if table_line.startswith('|') and table_line.endswith('|'):
                        # Remove first and last | and split
                        cells = table_line[1:-1].split('|')
                        table_data.append([cell.strip() for cell in cells])
                
                if table_data:
                    # Create table
                    num_rows = len(table_data)
                    num_cols = len(table_data[0])
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    # Fill table
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # Make header row bold
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
            continue
        
        # Handle mathematical equations
        if '$$' in line:
            # Extract equation
            equation_match = re.search(r'\$\$(.*?)\$\$', line)
            if equation_match:
                equation = equation_match.group(1)
                # Clean up equation
                equation = equation.replace('\\', '\\\\')
                p = doc.add_paragraph(f"Equation: {equation}", style='CustomEquation')
                i += 1
                continue
        
        # Handle equation explanations
        if line.startswith('**This equation') or line.startswith('**Where:') or line.startswith('**Symbol Meanings:') or line.startswith('**Range:'):
            p = doc.add_paragraph(line, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(bullet_text, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\. ', line):
            numbered_text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(numbered_text, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # Handle bold text
        if '**' in line:
            # Replace markdown bold with Word bold
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph(line, style='CustomNormal')
            # Make the paragraph bold
            for run in p.runs:
                run.bold = True
            i += 1
            continue
        
        # Handle regular text
        p = doc.add_paragraph(line, style='CustomNormal')
        i += 1
    
    # Save the document
    doc.save(output_file)
    print(f"Advanced Word document created successfully: {output_file}")

if __name__ == "__main__":
    input_file = "project report.md"
    output_file = "WSL_Project_Report_Advanced.docx"
    
    try:
        create_advanced_word_document(input_file, output_file)
        print("✅ Advanced conversion completed successfully!")
        print(f"📄 Advanced Word document saved as: {output_file}")
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}") 